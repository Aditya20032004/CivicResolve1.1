import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches


def add_title_page(doc: Document):
    doc.add_heading('CivicResolve v1.1 – Comprehensive Project Report', level=0)
    p = doc.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run(datetime.now().strftime('%Y-%m-%d'))
    p = doc.add_paragraph()
    p.add_run('Project: ').bold = True
    p.add_run('CivicResolvev1.1 – AI-Powered Civic Issue Management System')
    doc.add_page_break()


def add_overview_section(doc: Document):
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'CivicResolve v1.1 is an AI-powered civic issue management system that detects, '
        'validates, routes, and tracks urban infrastructure problems such as potholes, '
        'garbage, vandalism, dead animals, damaged roads, and more. It combines a YOLOv8 '
        'object detection model, a Flask backend, and a Next.js/React (TypeScript) frontend '
        'to deliver an end-to-end workflow from detection to verification.'
    )
    doc.add_paragraph(
        'The current version supports all 10 classes from the civic YOLO dataset and stores '
        'each class in its own database table, with automatic worker dispatch and anti-fraud '
        'validation on citizen uploads.'
    )

    doc.add_heading('1.1 Key Features', level=2)
    features = [
        'End-to-end pipeline from image capture to verified resolution and closure of incidents.',
        'YOLOv8-based multi-class detection across 10 real-world civic issue categories.',
        'Per-class database tables and admin views for fine-grained monitoring and reporting.',
        'Smart validator combining EXIF, GPS, screenshot heuristics, and model outputs to fight spam and fraud.',
        'Automatic worker dispatch using geospatial distance and workload balancing.',
        'Dual verification that cross-checks worker uploads against both citizen images and camera frames.',
        'Next.js/React frontends tailored for citizens, admins, workers, and camera operators.',
        'Configurable thresholds for validation scores, deduplication radius, and worker selection policies.',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('1.2 Primary Use Cases', level=2)
    use_cases = [
        'Continuous monitoring of key road segments in a city using fixed cameras.',
        'Fast reporting of potholes, garbage, and vandalism incidents by citizens via mobile or web.',
        'Prioritizing high-severity or high-traffic incidents for faster dispatch and resolution.',
        'Auditing contractor performance by comparing closure photos against original reports.',
        'Generating periodic reports for city authorities on incident trends and SLA compliance.',
    ]
    for u in use_cases:
        doc.add_paragraph(u, style='List Bullet')

    doc.add_heading('1.3 Background & Related Work', level=2)
    doc.add_paragraph(
        'Vision-based civic issue monitoring builds on a decade of progress in deep learning '
        'for object detection. Modern architectures such as the YOLO family achieve real-time '
        'performance on edge GPUs while maintaining strong accuracy on benchmarks like COCO. '
        'CivicResolve adapts these advances from generic benchmarks to a focused 10-class civic '
        'dataset, emphasizing real-world issues like potholes, vandalism, and garbage that '
        'directly affect citizens.'
    )
    doc.add_paragraph(
        'In parallel, many smart-city pilots have explored camera networks for traffic '
        'management, safety, and environmental monitoring. However, those systems are often '
        'closed, vendor-specific, and tightly coupled to particular hardware. CivicResolve takes '
        'a more open, developer-friendly approach: using Python, open-source detection models, '
        'and a transparent REST API so that cities, researchers, and startups can experiment '
        'rapidly without being locked into proprietary stacks.'
    )
    doc.add_paragraph(
        'From a workflow perspective, the project also draws ideas from modern incident '
        'management tools used in IT and SRE: tickets, SLAs, verification steps, and audit '
        'trails. By merging these ideas with computer vision, CivicResolve moves beyond pure '
        'detection and into a full life-cycle view of urban issues—from first detection to '
        'verified resolution and long-term analytics.'
    )


def parse_dataset_info(log_path: str) -> dict:
    """Parse dataset statistics (image counts, splits, classes) from project.log."""
    info: dict[str, int | None] = {
        "total_images": None,
        "valid_samples": None,
        "train_samples": None,
        "val_samples": None,
        "num_classes": None,
    }

    if not os.path.exists(log_path):
        return info

    with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            if "Found" in line and "total images" in line:
                m = re.search(r"Found (\d+) total images", line)
                if m:
                    info["total_images"] = int(m.group(1))
            elif "Found" in line and "valid samples" in line:
                m = re.search(r"Found (\d+) valid samples", line)
                if m:
                    info["valid_samples"] = int(m.group(1))
            elif "Saving" in line and "samples to train" in line:
                m = re.search(r"Saving (\d+) samples to train", line)
                if m:
                    info["train_samples"] = int(m.group(1))
            elif "Saving" in line and "samples to val" in line:
                m = re.search(r"Saving (\d+) samples to val", line)
                if m:
                    info["val_samples"] = int(m.group(1))
            elif "Config created" in line and "with" in line and "classes" in line:
                m = re.search(r"with (\d+) classes", line)
                if m:
                    info["num_classes"] = int(m.group(1))

    return info


def parse_training_runs(log_path: str) -> list[dict]:
    """Parse overall and per-class metrics for each completed training run.

    Returns a list of dicts; the last element corresponds to the most recent run.
    """
    runs: list[dict] = []
    if not os.path.exists(log_path):
        return runs

    current: dict | None = None
    in_per_class = False

    with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
        for line in f:
            if "TRAINING RESULTS & METRICS" in line:
                # Start a new run block
                if current:
                    runs.append(current)
                current = {
                    "timestamp": line.split(" - ")[0].strip(),
                    "mAP50": None,
                    "mAP50_95": None,
                    "precision": None,
                    "recall": None,
                    "val_mAP50": None,
                    "val_mAP50_95": None,
                    "val_precision": None,
                    "val_recall": None,
                    "per_class": {},
                }
                in_per_class = False
                continue

            if not current:
                continue

            # Overall training metrics
            if "mAP50:" in line and "Validation" not in line:
                m = re.search(r"mAP50: ([0-9.]+)", line)
                if m:
                    current["mAP50"] = float(m.group(1))
            elif "mAP50-95:" in line and "Validation" not in line:
                m = re.search(r"mAP50-95: ([0-9.]+)", line)
                if m:
                    current["mAP50_95"] = float(m.group(1))
            elif "Precision:" in line and "Validation" not in line:
                m = re.search(r"Precision: ([0-9.]+)", line)
                if m:
                    current["precision"] = float(m.group(1))
            elif "Recall:" in line and "Validation" not in line:
                m = re.search(r"Recall: ([0-9.]+)", line)
                if m:
                    current["recall"] = float(m.group(1))

            # Validation metrics
            if "Validation mAP50:" in line:
                m = re.search(r"Validation mAP50: ([0-9.]+)", line)
                if m:
                    current["val_mAP50"] = float(m.group(1))
            elif "Validation mAP50-95:" in line:
                m = re.search(r"Validation mAP50-95: ([0-9.]+)", line)
                if m:
                    current["val_mAP50_95"] = float(m.group(1))
            elif "Validation Precision:" in line:
                m = re.search(r"Validation Precision: ([0-9.]+)", line)
                if m:
                    current["val_precision"] = float(m.group(1))
            elif "Validation Recall:" in line:
                m = re.search(r"Validation Recall: ([0-9.]+)", line)
                if m:
                    current["val_recall"] = float(m.group(1))

            # Per-class block
            if "Per-Class mAP50-95:" in line:
                in_per_class = True
                continue

            if in_per_class:
                # Stop if we hit a separator or empty line
                if line.strip() == "" or "====" in line or "GENERATING VISUALIZATIONS" in line:
                    in_per_class = False
                    continue
                # Example: "  Vandalism Issues: 0.5194"
                m = re.search(r"INFO -\s+(.+?):\s+([0-9.]+)", line)
                if m:
                    cls_name = m.group(1).strip()
                    value = float(m.group(2))
                    current["per_class"][cls_name] = value

    if current:
        runs.append(current)
    return runs


def add_architecture_section(doc: Document):
    doc.add_heading('2. System Architecture', level=1)

    doc.add_heading('2.1 Components', level=2)
    components = [
        ('Frontend (Next.js / React)', 'Citizen, admin, and worker UI in frontend_4.'),
        ('Backend (Flask API)', 'REST endpoints for citizen reports, admin stats, workflow control, and AI helpers.'),
        ('AI/ML Layer', 'YOLOv8 model and training utilities under ai_ml.'),
        ('Workflow Layer', 'Python workflows to orchestrate detection, severity, worker assignment, and verification.'),
        ('Camera Integration', 'Laptop/camera feed integration and nearest-frame lookup.'),
        ('Persistence', 'SQLite database and on-disk image storage under data/images.'),
    ]
    for name, desc in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}: ').bold = True
        p.add_run(desc)

    doc.add_heading('2.2 High-Level Flow', level=2)
    doc.add_paragraph(
        '1. City cameras or a laptop webcam stream frames into the camera integration service, '
        'which runs YOLOv8 to detect civic issues in real time.'
    )
    doc.add_paragraph(
        '2. Citizens can also submit photos and metadata (type, latitude, longitude, address) '
        'via the web UI, which calls the /api/citizen/report endpoint.'
    )
    doc.add_paragraph(
        '3. The backend validates the image using EXIF/GPS checks, screenshot heuristics, and '
        'YOLO content verification, then normalizes the issue type and performs spatial '
        'deduplication per class.'
    )
    doc.add_paragraph(
        '4. New incidents are stored in per-class tables and automatically assigned to the best '
        'available worker based on proximity and workload.'
    )
    doc.add_paragraph(
        '5. Workers upload resolved images; the verification workflow cross-checks them against '
        'the original report and the nearest camera frame before approving or rejecting the fix.'
    )
    doc.add_paragraph(
        '6. Admin dashboards aggregate incidents across all classes, track SLA-style metrics, and '
        'provide verification and camera sweep controls.'
    )

    doc.add_paragraph(
        'The architecture intentionally keeps the AI/ML layer, HTTP API, and UI concerns loosely '
        'coupled. The YOLOv8 model is wrapped behind thin Python utilities, while the Flask '
        'backend exposes stable JSON APIs that the Next.js frontend consumes. This separation '
        'allows the model to be upgraded independently (for example, retraining on additional '
        'cities or adding new classes) without breaking the user-facing workflows.'
    )
    doc.add_paragraph(
        'Persistence is handled through SQLite for simplicity, but the ORM-based models make it '
        'straightforward to move to PostgreSQL or a managed cloud database. The same pattern '
        'applies to deployment: the project can run locally for experimentation or be deployed '
        'to cloud platforms with minimal configuration changes.'
    )

    doc.add_paragraph()
    doc.add_paragraph('Figure 1: High-level architecture diagram of CivicResolve v1.1 (insert diagram here).').italic = True


def add_classes_table(doc: Document):
    doc.add_heading('3. Civic Classes and Database Design', level=1)
    doc.add_paragraph(
        'CivicResolve uses 10 classes from the civic YOLO dataset. Each class is normalized to a '
        'canonical internal type and mapped to a dedicated SQLAlchemy model and database table.'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Canonical Type'
    hdr_cells[1].text = 'Example Raw Labels'
    hdr_cells[2].text = 'Model Class'
    hdr_cells[3].text = 'DB Table'

    rows = [
        ('damaged_road', '"Damaged Road issues", "road damage"', 'DamagedRoadReport', 'damaged_roads'),
        ('pothole', '"Pothole Issues", "potholes"', 'PotholeReport', 'potholes'),
        ('illegal_parking', '"Illegal Parking Issues"', 'IllegalParkingReport', 'illegal_parking'),
        ('broken_sign', '"Broken Road Sign Issues", "broken sign"', 'BrokenSignReport', 'broken_signs'),
        ('fallen_tree', '"Fallen trees"', 'FallenTreeReport', 'fallen_trees'),
        ('garbage', '"Littering/Garbage on Public Places", "garbage", "trash"', 'GarbageReport', 'garbage'),
        ('vandalism', '"Vandalism Issues"', 'VandalismReport', 'vandalism'),
        ('dead_animal', '"Dead Animal Pollution"', 'DeadAnimalReport', 'dead_animals'),
        ('damaged_concrete', '"Damaged concrete structures"', 'DamagedConcreteReport', 'damaged_concrete'),
        ('damaged_wires', '"Damaged Electric wires and poles"', 'DamagedWiresReport', 'damaged_wires'),
    ]

    for canonical, labels, model, table_name in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = canonical
        row_cells[1].text = labels
        row_cells[2].text = model
        row_cells[3].text = table_name

    doc.add_paragraph()
    doc.add_paragraph('Table 1: Mapping from YOLO civic classes to internal models and database tables.').italic = True


def add_validation_and_workflows(doc: Document):
    doc.add_heading('4. Validation, Severity, and Workflows', level=1)

    doc.add_heading('4.1 Smart Issue Validator', level=2)
    doc.add_paragraph(
        'The IssueValidator combines EXIF metadata analysis, GPS consistency checks, '
        'screenshot detection, and YOLO content verification to compute a trust score between '
        '0 and 100. Based on this score, the report is classified as approved, flagged, or '
        'rejected before entering the main workflow.'
    )
    doc.add_paragraph(
        'EXIF analysis looks for missing or suspicious GPS coordinates, time anomalies, and '
        'device signatures. Screenshot detection uses basic heuristics on resolution patterns '
        'and color histograms to catch images captured from screens or social media. The YOLO '
        'content check verifies that at least one detection matches the claimed issue type, '
        'preventing users from mislabeling images to game the system.'
    )

    doc.add_heading('4.2 Severity Estimation', level=2)
    doc.add_paragraph(
        'A severity classifier aggregates bounding-box areas per class and estimates severity '
        'levels (low, medium, high) based on the proportion of the image occupied by the '
        'dominant issue or, if unavailable, by the count of detections.'
    )
    doc.add_paragraph(
        'For example, a pothole that occupies more than one-third of the road width in an image '
        'is treated as high severity, while scattered garbage occupying a small fraction of the '
        'frame may be treated as low or medium. These thresholds are configurable and can be '
        'tuned per city based on field feedback from engineers and workers.'
    )

    doc.add_heading('4.3 Worker Assignment Workflow', level=2)
    doc.add_paragraph(
        'The worker assignment workflow maintains a pool of 20 Bhopal-based workers. It selects '
        'the nearest eligible worker for each new incident using a haversine distance metric, '
        'current active tasks, and a simple reliability score, then marks the incident as '
        'assigned.'
    )
    doc.add_paragraph(
        'The workflow also ensures that the same worker is not overloaded and that incidents '
        'are spread across the pool. This makes the system fair while still prioritizing '
        'geographical proximity to minimize travel time. In future versions, this can be '
        'extended with richer constraints such as shift timings and skill specialization.'
    )

    doc.add_heading('4.4 Dual Verification Workflow', level=2)
    doc.add_paragraph(
        'After a worker uploads a resolved image, the dual verification workflow compares the '
        'resolved photo with both the original report image and the latest nearby camera frame. '
        'It uses YOLO confidences and color-histogram similarity to decide whether the issue '
        'has truly been resolved before updating the incident status.'
    )
    doc.add_paragraph(
        'This design intentionally does not trust any single source of truth. A worker cannot '
        'simply upload a random clean road photo; the system cross-checks the spatial context '
        'and content against independently captured camera footage. This reduces the risk of '
        'false closures and helps build trust in the end-to-end pipeline for both citizens and '
        'administrators.'
    )


def add_training_metrics_section(doc: Document, dataset_info: dict, latest_run: dict | None):
    doc.add_heading('5. Training Dataset & Model Performance', level=1)

    # Dataset table
    doc.add_heading('5.1 Dataset Summary', level=2)
    p = doc.add_paragraph(
        'The YOLO training dataset is prepared from the raw civic archive under data/archive. '
        'The data preparation step scans all images, filters out samples without labels, and '
        'creates a YOLO-formatted dataset.yaml with 10 classes.'
    )
    p

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Value'

    def add_row(label: str, value: str):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    if dataset_info:
        add_row('Total images scanned', str(dataset_info.get('total_images') or 'n/a'))
        add_row('Valid labeled samples', str(dataset_info.get('valid_samples') or 'n/a'))
        add_row('Train samples', str(dataset_info.get('train_samples') or 'n/a'))
        add_row('Validation samples', str(dataset_info.get('val_samples') or 'n/a'))
        add_row('Number of classes', str(dataset_info.get('num_classes') or 'n/a'))
    else:
        add_row('Dataset statistics', 'Not available (project.log not found)')

    doc.add_paragraph()
    doc.add_paragraph('Table 2: Summary of the YOLO training dataset used for CivicResolve v1.1.').italic = True

    # Training metrics
    doc.add_heading('5.2 Final Training & Validation Metrics', level=2)
    if not latest_run:
        doc.add_paragraph('Training metrics were not found in logs/project.log.').italic = True
        return

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Metric'
    hdr2[1].text = 'Train'
    hdr2[2].text = 'Validation'

    metrics = [
        ('mAP50', latest_run.get('mAP50'), latest_run.get('val_mAP50')),
        ('mAP50-95', latest_run.get('mAP50_95'), latest_run.get('val_mAP50_95')),
        ('Precision', latest_run.get('precision'), latest_run.get('val_precision')),
        ('Recall', latest_run.get('recall'), latest_run.get('val_recall')),
    ]

    for name, train_v, val_v in metrics:
        row = table2.add_row().cells
        row[0].text = name
        row[1].text = f"{train_v:.4f}" if isinstance(train_v, float) else 'n/a'
        row[2].text = f"{val_v:.4f}" if isinstance(val_v, float) else 'n/a'

    doc.add_paragraph()
    doc.add_paragraph('Table 3: Final training and validation metrics for the latest YOLO model run.').italic = True

    # Per-class metrics table if available
    if latest_run.get('per_class'):
        doc.add_heading('5.3 Per-Class mAP50-95', level=2)
        per_class = latest_run['per_class']
        table3 = doc.add_table(rows=1, cols=2)
        table3.style = 'Light Grid Accent 1'
        hdr3 = table3.rows[0].cells
        hdr3[0].text = 'Class'
        hdr3[1].text = 'mAP50-95'

        for cls_name, value in per_class.items():
            row = table3.add_row().cells
            row[0].text = cls_name
            row[1].text = f"{value:.4f}"

        doc.add_paragraph()
        doc.add_paragraph('Table 4: Per-class validation mAP50-95 for the latest YOLO model.').italic = True

    # Embed existing visualizations if present
    doc.add_heading('5.4 Training Curves and Confusion Matrix', level=2)
    vis_dir = os.path.join('ai_ml', 'runs', 'visualizations')
    if os.path.isdir(vis_dir):
        training_curves = os.path.join(vis_dir, 'training_curves.png')
        confusion = os.path.join(vis_dir, 'confusion_matrix.png')
        val_preds = os.path.join(vis_dir, 'validation_predictions.png')

        if os.path.exists(training_curves):
            try:
                doc.add_picture(training_curves, width=Inches(5.5))
                doc.add_paragraph('Figure 2: YOLO training and validation curves (loss/metrics over epochs).').italic = True
            except Exception:
                doc.add_paragraph('Figure 2: [Training curves could not be embedded.]').italic = True
        else:
            doc.add_paragraph('Figure 2: [Training curves image not found – expected at ai_ml/runs/visualizations/training_curves.png.]').italic = True

        if os.path.exists(confusion):
            try:
                doc.add_picture(confusion, width=Inches(5.5))
                doc.add_paragraph('Figure 3: Confusion matrix for the 10-class civic YOLO model.').italic = True
            except Exception:
                doc.add_paragraph('Figure 3: [Confusion matrix could not be embedded.]').italic = True
        else:
            doc.add_paragraph('Figure 3: [Confusion matrix image not found – expected at ai_ml/runs/visualizations/confusion_matrix.png.]').italic = True

        if os.path.exists(val_preds):
            try:
                doc.add_picture(val_preds, width=Inches(5.5))
                doc.add_paragraph('Figure 4: Sample validation predictions with bounding boxes and class labels.').italic = True
            except Exception:
                doc.add_paragraph('Figure 4: [Validation predictions image could not be embedded.]').italic = True
        else:
            doc.add_paragraph('Figure 4: [Validation predictions image not found – expected at ai_ml/runs/visualizations/validation_predictions.png.]').italic = True
    else:
        doc.add_paragraph('Figures 2–4: [Visualization directory ai_ml/runs/visualizations not found.]').italic = True


def add_evaluation_section(doc: Document):
    doc.add_heading('6. Evaluation & Sample Results', level=1)

    doc.add_heading('6.1 Sample Vandalism Detection', level=2)
    doc.add_paragraph(
        'On a sample frame from the Bhopal Node camera, the YOLOv8 model detected 8 regions '
        'labeled as "Vandalism Issues" at 640×640 resolution. Detection confidences for this '
        'frame ranged approximately from 0.26 to 0.85, confirming that the model consistently '
        'identifies vandalism in the test scene. The annotated image is saved under data/images '
        'with a filename such as Vandalism_Issues_...jpg.'
    )

    # Optional: try to embed one example image if it exists
    sample_img_dir = os.path.join('data', 'images')
    sample_img = None
    if os.path.isdir(sample_img_dir):
        for name in os.listdir(sample_img_dir):
            if name.lower().endswith(('.jpg', '.jpeg', '.png')) and 'vandalism' in name.lower():
                sample_img = os.path.join(sample_img_dir, name)
                break

    if sample_img and os.path.exists(sample_img):
        try:
            doc.add_picture(sample_img, width=Inches(4.5))
            doc.add_paragraph('Figure 5: Sample YOLO vandalism detection on a Bhopal Node frame.').italic = True
        except Exception:
            doc.add_paragraph('Figure 5: [Image not embedded – error loading example vandalism image.]').italic = True
    else:
        doc.add_paragraph('Figure 5: [Placeholder – insert sample vandalism detection image here.]').italic = True

    doc.add_heading('6.2 Validator Decision Example', level=2)
    doc.add_paragraph(
        'For the same upload, the smart validator produced a trust score of 85 and a decision '
        'of "approved", with EXIF missing (no GPS data), screenshot not detected, and content '
        'status "match" for the claimed type "Vandalism Issues". Since the score exceeded the '
        'approval threshold (80), the report proceeded into the main workflow.'
    )

    # Summary metrics table (can be expanded with real aggregated stats later)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Sample Value / Observation'

    metrics = [
        ('Detection latency per frame (640x640)', '~20 ms'),
        ('Number of vandalism detections (sample)', '8'),
        ('Validator score (sample)', '85'),
        ('Validator decision (sample)', 'approved'),
        ('Deduplication radius', '~0.08 km (configurable)'),
        ('Worker auto-assignment', 'Enabled via proximity + workload heuristic'),
    ]

    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value

    doc.add_paragraph()
    doc.add_paragraph('Table 5: Sample evaluation metrics for vandalism detection and validation.').italic = True

    doc.add_heading('6.3 Graphs and Longitudinal Evaluation', level=2)
    doc.add_paragraph(
        'Beyond single-sample analysis, the system is designed to support longitudinal '
        'evaluation over weeks or months. By aggregating incidents per class, tracking mean '
        'time-to-acknowledge and time-to-resolution, and monitoring validator score '
        'distributions, city operators can study whether interventions are actually improving '
        'conditions on the ground.'
    )
    doc.add_paragraph(
        'For example, a reduction in vandalism incidents around a newly installed CCTV camera '
        'cluster, or a drop in garbage complaints after a route optimization for collection '
        'vehicles, would show up directly in these graphs. Such insights turn raw detections '
        'into actionable policy feedback loops.'
    )
    doc.add_paragraph('Figure 6: [Placeholder – insert incident counts per class bar chart here.]').italic = True


def add_conclusion(doc: Document):
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        'CivicResolve v1.1 delivers an end-to-end pipeline for AI-assisted civic issue '
        'management: YOLO-based detection, smart validation, per-class storage, worker '
        'assignment, and dual verification with camera cross-checks. The modular design across '
        'backend, frontend, AI, and workflows makes the system easy to extend to new cities, '
        'issue classes, and deployment environments.'
    )
    doc.add_paragraph(
        'In practical terms, this means that a single platform can support citizens reporting '
        'issues from their phones, municipal staff monitoring live camera feeds, and on-ground '
        'workers closing tickets with strong evidence that the underlying problem was actually '
        'fixed. By structuring incidents per class and exposing rich APIs, the project also '
        'lays the groundwork for analytics dashboards and research on urban infrastructure '
        'quality over time.'
    )
    doc.add_paragraph(
        'Future work could explore more advanced model architectures (for example, segmenting '
        'damaged regions for better quantification), integrating route optimization for worker '
        'scheduling, and adding citizen feedback loops after resolution. Another direction is to '
        'deploy the same stack across multiple cities and compare how incident patterns, '
        'response times, and vandalism rates differ, turning CivicResolve into a broader civic '
        'analytics platform rather than just a ticketing system.'
    )


def main():
    os.makedirs('docs', exist_ok=True)

    doc = Document()
    add_title_page(doc)
    add_overview_section(doc)
    add_architecture_section(doc)
    add_classes_table(doc)
    add_validation_and_workflows(doc)

    log_path = os.path.join('logs', 'project.log')
    dataset_info = parse_dataset_info(log_path)
    runs = parse_training_runs(log_path)
    latest_run = runs[-1] if runs else None

    add_training_metrics_section(doc, dataset_info, latest_run)
    add_evaluation_section(doc)
    add_conclusion(doc)

    output_path = os.path.join('docs', 'CivicResolve_v1.1_Report.docx')
    doc.save(output_path)
    print(f'Report generated at: {output_path}')


if __name__ == '__main__':
    main()
